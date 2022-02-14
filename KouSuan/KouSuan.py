# !/usr/bin/env python
# -*- coding: utf-8 -*-
# @Author: gyy
import random
from docx import Document
import openpyxl
from openpyxl.utils.dataframe import dataframe_to_rows
import pandas as pd
from alive_progress import alive_bar
import os, re, gc

class KouSuan:
    def __init__(self, path):
        self.ques_list = []
        self.operators = [">", "<", "=", "≈", "-", "+", "×", "÷", "/"]
        self.path = path
        self.excel = self.find_file(".xlsx", path)

    def find_file(self, type, path, choose=True):
        file_list = []
        for root, dirs, files in os.walk(path):
            for f in files:
                if type in f:
                    file_list.append(f)
                    if not choose:
                        return f
        if choose:
            print("检测到以下{}文件: ".format(type))
            for f in file_list:
                print("{}: {}".format(file_list.index(f), f))
            while True:
                try:
                    target = file_list[int(input("请输入文件序号："))]
                    return target
                except:
                    print("发生错误：请正确输入文件前的序号（0-n）")

    def generate_option(self, n):
        def roll(num, m=1, is_float=False):
            ls = []
            ls.append(num)
            for i in range(3):
                j = 0
                while True:
                    if is_float:
                        option = random.uniform(num / 2, num + num / 2 + j)
                        option = round(option, m)
                    else:
                        option = random.randrange(num // 2, num + num // 2 + j)
                    if option not in ls and option != 0:
                        ls.append(option)
                        break
                    else:
                        j += 1
            return ls

        ans_list = []
        if n in self.operators:
            if n in ["<", ">", "="]:
                ans_list = ["<", ">", "=", "无法比较"]
            else:
                ans_list = ["+", "-", "×", "÷"]
        else:
            try:
                n = re.sub('x=|%|℃', '', n).strip()
                if '/' in n:
                    a, b = re.findall('\d+', n)
                    a_ls = roll(int(a))
                    for i in range(4):
                        if '(' in n:
                            ans = '(' + str(a_ls[i]) + '/' + str(b) + ')'
                        else:
                            ans = str(a_ls[i]) + '/' + str(b)
                        ans_list.append(ans)
                    random.shuffle(ans_list)
                elif '.' in n:
                    m = len(n.split('.')[-1])
                    n_f = float(n)
                    ans_list = roll(n_f, m, is_float=True)
                    random.shuffle(ans_list)
                else:
                    n = int(n)
                    ans_list.append(n)
                    if n % 10 != 0 and n != 10:
                        ans_list = roll(n)
                    else:
                        for i in range(3):
                            j = 0
                            while True:
                                m = random.randrange(-2, 4)
                                option = abs(n + (m + j) * 10)
                                if option not in ans_list:
                                    ans_list.append(option)
                                    break
                                else:
                                    j += 1
                    random.shuffle(ans_list)
            except Exception as e:
                return ""
        ans = "|".join([str(i) for i in ans_list])
        return ans

    def get_element(self, p_list):
        def empty():
            d = {}
            d["num1"] = ""
            d["operator1"] = ""
            d["num2"] = ""
            d["operator2"] = ""
            d["question"] = ""
            d["final_ans"] = "wrong"
            return d
        for i in p_list:
            info_dict = {}
            res = re.split("(\+|＋|-|×|÷|/|□|£|=|≈)", i.replace(" ", ""))
            res = list(filter(None, res))
            # print(i)
            if len(res) < 2:
                self.ques_list.append(empty())
                continue
            info_dict["num1"] = res[0]
            info_dict["operator1"] = res[1]
            if "=" in res or "≈" in res:
                equal_loc = ""
                if "=" in res:
                    equal_loc = res.index("=")
                    info_dict["operator2"] = "="
                elif "≈" in res:
                    equal_loc = res.index("≈")
                    info_dict["operator2"] = "≈"
                info_dict["num2"] = "".join(res[2:equal_loc])
                if "□" not in res and "（）" not in res:
                    info_dict["question"] = "mathresult"
                    info_dict["final_ans"] = "answer"
                else:
                    space_loc = 0
                    move = 1
                    if "□" in res:
                        space_loc = res.index("□")
                        # move = 1
                    if "（）" in res:
                        space_loc = res.index("（）")
                        # move = 2
                    if space_loc == 0:
                        info_dict["question"] = "num1"
                    elif space_loc == 1:
                        info_dict["question"] = "mathematic"
                    else:
                        info_dict["question"] = "num2"
                    info_dict["final_ans"] = "".join(res[equal_loc + move:])
            # 比较
            elif "□" in res:
                space_loc = res.index("□")
                info_dict["num2"] = "".join(res[2:space_loc])
                info_dict["operator2"] = ""
                info_dict["question"] = "boolval"
                info_dict["final_ans"] = "".join(res[space_loc + 1:])
            else:
                self.ques_list.append(empty())
                continue
            self.ques_list.append(info_dict)
            # if info_dict['question'] == "num2":
            #     print(info_dict)

    def word2excel(self, excel, has_ques=False):
        p_list = []
        if has_ques:
            df = pd.read_excel(self.path + excel)
            p_list = df['习题问题文本'].fillna("").apply(lambda x: str(x).replace("卍", "")).tolist()
        else:
            word = self.find_file(".docx", self.path)
            doc = Document(self.path + word)
            if len(doc.paragraphs) > 1:
                for i in doc.paragraphs:
                    p_list.append(i.text.replace("\n", "").replace("卍", ""))
            else:
                p_list = doc.paragraphs[0].text.replace("卍", "").split('\n')

        self.get_element(p_list)
        # print(self.ques_list)
        wb = openpyxl.load_workbook(self.path + excel, data_only=True)
        ws = wb.active

        ws.cell(1, 12).value = "卡片配置-自动模式"
        ws.cell(1, 13).value = "卡片配置-问题文本"
        ws.cell(1, 14).value = "习题问题文本"
        ws.cell(1, 15).value = "标题14"
        ws.cell(1, 16).value = "计算数1"
        ws.cell(1, 17).value = "运算符"
        ws.cell(1, 18).value = "计算数2"
        ws.cell(1, 19).value = "表达式"
        ws.cell(1, 20).value = "答案"
        ws.cell(1, 21).value = "哪里问号"

        length = len(self.ques_list)
        del_rows = []
        with alive_bar(length) as bar:
            for i in range(length):
                row = i + 2
                col = 10
                answer = str(ws.cell(row, col).value)
                option = self.generate_option(answer)
                if option == "":
                    continue
                info_dict = self.ques_list[i]

                if answer in ["<", ">", "=", "≈"]:
                    info_dict["operator2"] = answer
                for key, value in info_dict.items():
                    if "□" in value:
                        info_dict[key] = value.replace("□", answer)
                    if "（）" in value:
                        info_dict[key] = value.replace("（）", answer)

                ws.cell(row, col + 4).value = p_list[i]
                ws.cell(row, col + 6).value = info_dict["num1"]
                ws.cell(row, col + 7).value = info_dict["operator1"]
                ws.cell(row, col + 8).value = info_dict["num2"]
                ws.cell(row, col + 9).value = info_dict["operator2"]
                ws.cell(row, col + 11).value = info_dict["question"]
                ws.cell(row, col + 1).value = option

                final_ans = info_dict["final_ans"]
                if final_ans:
                    if final_ans == "wrong":
                        del_rows.append(row)
                        ws.cell(row, col + 10).value = "wrong"
                    elif final_ans == "answer":
                        ws.cell(row, col + 10).value = answer
                    else:
                        ws.cell(row, col + 10).value = final_ans
                bar()
        wb.save(self.path + "res-" + excel)
        del wb, ws
        gc.collect()

    def sample(self, n):
        excel = self.find_file(".xlsx", self.path, choose=True)
        df = pd.read_excel(self.path + "res-" + excel, dtype=str)
        df.dropna(subset=[df.columns[10], '计算数1', '运算符', '计算数2', '表达式', '答案'], inplace=True)
        new_df = pd.DataFrame()
        option_list = list(set(df.iloc[:, 20].dropna().tolist()))
        for i in option_list:
            temp = df[df.iloc[:, 20] == i]
            if temp.size > n:
                temp = temp.iloc[:n, :]
            if option_list.index(i) == 0:
                new_df = temp
            else:
                new_df = pd.concat([new_df, temp])
        # print(df.head()['表达式'])
        self.openpyxl_saver(df, self.path + "res-" + excel)
        self.openpyxl_saver(new_df, self.path + "sample-" + excel)

    def openpyxl_saver(self, df, name):
        wb = openpyxl.Workbook()
        ws = wb.active
        for r in dataframe_to_rows(df, index=False):
            ws.append(r)
        wb.save(name)
        del wb, ws
        gc.collect()

    def run(self):
        self.word2excel(self.excel, has_ques=True)
        self.sample(20)

def main():
    path = r"口算属性表/"
    ks = KouSuan(path)
    ks.run()

main()
